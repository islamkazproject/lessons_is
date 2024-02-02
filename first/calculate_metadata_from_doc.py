"""
Задание 1.

Разработайте функцию Python, которая читает doc файл
и вычисляет метаданные, такие как: количество строк, слов и символов в файле.
Требуется игнорировать пробелы и знаки препинания при подсчете символов.
В эту программу будет только один вход: doc файл.
После обработки результат будет выведен на экран.

"""
import re
import docx


def calculate_count_paragraph(document, with_empty=True):
    paragraph_count = 0
    for paragraph in document.paragraphs:
        if with_empty or paragraph.text.strip():
            paragraph_count += 1

    return paragraph_count


def calculate_count_word(document):
    words_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        words = re.findall(r"\w+", text)
        words_count += len(words)

    return words_count


def calculate_count_symbol(document):
    symbols_count = 0
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            exclude_symbols = r"\W" + "".join([".", ",", "?", "!", ":", ";", "\"", "(", ")", "-", "–"])
            symbols_count += len(re.sub(exclude_symbols, "", run.text))
    return symbols_count


def count_metadata(doc_path: str):
    document = docx.Document(doc_path)

    paragraph_count_without_empty = calculate_count_paragraph(document, with_empty=False)
    paragraph_count_with_empty = calculate_count_paragraph(document, with_empty=True)
    count_word = calculate_count_word(document)
    count_symbol = calculate_count_symbol(document)

    return paragraph_count_without_empty, paragraph_count_with_empty, count_word, count_symbol


if __name__ == '__main__':
    doc_path = "/Users/islam/PycharmProjects/lessons_is/first/text.docx"
    paragraph_count_without_empty, paragraph_count_with_empty, count_word, count_symbol = count_metadata(doc_path)
    print(f'Количество абзацев в файле без учета пустоты: {paragraph_count_without_empty}')
    print(f'Количество абзацев в файле с учетом пустоты: {paragraph_count_with_empty}')
    print(f'Количество слов в файле: {count_word}')
    print(f'Количество символов в файле: {count_symbol}')
